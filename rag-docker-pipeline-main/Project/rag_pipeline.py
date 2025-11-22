import os
import hashlib
import logging
import tempfile
import zipfile
import json
import csv
import re
import nltk
import docx
import whisper
import pdfplumber
import pytesseract
import cv2

from pydub import AudioSegment
from nltk.tokenize import sent_tokenize
from llama_index.core import VectorStoreIndex, Document, StorageContext, load_index_from_storage, Settings
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core.response_synthesizers import ResponseMode
from llama_index.llms.groq import Groq
from llama_index.llms.ollama import Ollama
from llama_index.core.node_parser import SentenceSplitter

nltk.download("punkt", quiet=True)
nltk.download("punkt_tab", quiet=True)

# === Config ===
PERSIST_DIR = "./storage"
HASH_CACHE_FILE = "text_hash.txt"
EMBEDDING_MODEL = "intfloat/e5-small-v2"
GROQ_MODEL = "llama-3.3-70b-versatile"
GROQ_API_KEY = "entrr api key here"  # Replace with your actual Groq API key
OLLAMA_MODEL = "gemma2:2b"
CHUNK_SIZE = 512
CHUNK_OVERLAP = 50

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# === Initialize Settings ===
def initialize_settings(use_groq=True):
    """Initialize LlamaIndex settings with Groq or Ollama"""
    try:
        # Set embedding model
        Settings.embed_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL)
        
        # Set chunk settings
        Settings.chunk_size = CHUNK_SIZE
        Settings.chunk_overlap = CHUNK_OVERLAP
        
        if use_groq:
            try:
                # Try Groq first
                if not GROQ_API_KEY or GROQ_API_KEY == "your_groq_api_key_here":
                    logging.warning("GROQ_API_KEY not configured. Please update GROQ_API_KEY in the script.")
                    raise ValueError("Missing GROQ_API_KEY")
                
                Settings.llm = Groq(
                    model=GROQ_MODEL,
                    api_key=GROQ_API_KEY,
                    temperature=0.1
                )
                logging.info(f"Using Groq LLM: {GROQ_MODEL}")
                return "groq"
            except Exception as e:
                logging.warning(f"Groq initialization failed: {e}. Falling back to Ollama...")
                raise
        
    except Exception as e:
        # Fallback to Ollama
        try:
            Settings.llm = Ollama(
                model=OLLAMA_MODEL,
                base_url="http://host.docker.internal:11434",
                temperature=0.1
            )
            logging.info(f"✅ Using Ollama LLM: {OLLAMA_MODEL}")
            return "ollama"
        except Exception as ollama_error:
            logging.error(f"Both Groq and Ollama failed: {ollama_error}")
            raise

# === File Extractors ===

def extract_text_from_pdf(path):
    try:
        with pdfplumber.open(path) as pdf:
            return "\n".join([p.extract_text() or "" for p in pdf.pages])
    except Exception as e:
        logging.error(f"PDF extraction failed: {e}")
        return ""

def extract_text_from_image(path):
    try:
        img = cv2.imread(path)
        return pytesseract.image_to_string(img)
    except Exception as e:
        logging.error(f"Image OCR failed: {e}")
        return ""

def extract_text_from_docx(path):
    try:
        return "\n".join([p.text for p in docx.Document(path).paragraphs])
    except Exception as e:
        logging.error(f"DOCX extraction failed: {e}")
        return ""

def extract_text_from_txt(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        logging.error(f"TXT extraction failed: {e}")
        return ""

def extract_text_from_csv(path):
    try:
        with open(path, newline='', encoding='utf-8') as csvfile:
            return "\n".join([" | ".join(row) for row in csv.reader(csvfile)])
    except Exception as e:
        logging.error(f"CSV extraction failed: {e}")
        return ""

def extract_text_from_json(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.dumps(json.load(f), indent=2)
    except Exception as e:
        logging.error(f"JSON extraction failed: {e}")
        return ""

def extract_text_from_audio(path):
    try:
        model = whisper.load_model("base")
        return model.transcribe(path)["text"]
    except Exception as e:
        logging.error(f"Audio transcription failed: {e}")
        return ""

def extract_text_from_zip(path):
    text = ""
    try:
        with zipfile.ZipFile(path, 'r') as zip_ref:
            with tempfile.TemporaryDirectory() as temp_dir:
                zip_ref.extractall(temp_dir)
                for root, _, files in os.walk(temp_dir):
                    for name in files:
                        full_path = os.path.join(root, name)
                        text += extract_text(full_path) + "\n"
    except Exception as e:
        logging.error(f"ZIP extraction failed: {e}")
    return text

def extract_text(path):
    """File dispatcher for text extraction"""
    ext = os.path.splitext(path)[1].lower()
    extractors = {
        ".pdf": extract_text_from_pdf,
        ".jpg": extract_text_from_image,
        ".jpeg": extract_text_from_image,
        ".png": extract_text_from_image,
        ".docx": extract_text_from_docx,
        ".txt": extract_text_from_txt,
        ".csv": extract_text_from_csv,
        ".json": extract_text_from_json,
        ".mp3": extract_text_from_audio,
        ".wav": extract_text_from_audio,
        ".mp4": extract_text_from_audio,
        ".zip": extract_text_from_zip
    }
    
    extractor = extractors.get(ext)
    if extractor:
        return extractor(path)
    else:
        logging.warning(f"Unsupported file type: {ext}")
        return ""

# === Utility Functions ===

def clean_text(text):
    """Clean and normalize text"""
    return re.sub(r'\s+', ' ', text).replace("•", "-").replace("–", "-").strip()

def compute_hash(text):
    """Compute MD5 hash of text"""
    return hashlib.md5(text.encode("utf-8")).hexdigest()

def has_file_changed(text):
    """Check if file content has changed"""
    if os.path.exists(HASH_CACHE_FILE):
        with open(HASH_CACHE_FILE, "r") as f:
            return compute_hash(text) != f.read().strip()
    return True

def cache_hash(text):
    """Cache the hash of current text"""
    with open(HASH_CACHE_FILE, "w") as f:
        f.write(compute_hash(text))

# === LlamaIndex RAG Functions ===

def create_index(file_paths):
    """Create or load vector index from files"""
    # Extract and combine text from all files
    documents = []
    for path in file_paths:
        text = clean_text(extract_text(path))
        if text:
            doc = Document(
                text=text,
                metadata={"source": os.path.basename(path)}
            )
            documents.append(doc)
    
    if not documents:
        logging.warning("No valid text extracted from files.")
        return None
    
    # Combine all text for hash checking
    combined_text = "\n".join([doc.text for doc in documents])
    
    # Check if we need to rebuild the index
    if os.path.exists(PERSIST_DIR) and not has_file_changed(combined_text):
        logging.info("📂 Loading existing index...")
        storage_context = StorageContext.from_defaults(persist_dir=PERSIST_DIR)
        index = load_index_from_storage(storage_context)
    else:
        logging.info("🔨 Creating new index...")
        index = VectorStoreIndex.from_documents(
            documents,
            show_progress=True
        )
        index.storage_context.persist(persist_dir=PERSIST_DIR)
        cache_hash(combined_text)
    
    return index

def query_index(index, query, llm_type):
    """Query the index with the given question"""
    try:
        # Create query engine with custom prompt
        query_engine = index.as_query_engine(
            similarity_top_k=5,
            response_mode=ResponseMode.TREE_SUMMARIZE,  # best formatting!
        )
        
        # Custom system prompt
        custom_prompt = (
            "You are a knowledge-based assistant designed to answer user queries exclusively using the provided documentation. "
            "Your responses should be drawn solely from the contents of the provided documentation. "
            "If a query cannot be answered based on the provided documentation, clearly state that the information is not available in the documents. "
            "Always ensure that your answers are consistent, factual, and directly relevant to the documentation.\n\n"
            f"Question: {query}"
        )
        
        response = query_engine.query(custom_prompt)
        
        # Format output
        source_info = f"🤖 {llm_type.upper()}" if llm_type else "📚 Documents"
        print(f"\n{source_info}")
        print(f"✅ Answer: {response.response.strip()}")
        
        # Show source documents if available
        if hasattr(response, 'source_nodes') and response.source_nodes:
            print("\n📄 Sources:")
            for idx, node in enumerate(response.source_nodes[:3], 1):
                source = node.metadata.get('source', 'Unknown')
                score = node.score if hasattr(node, 'score') else 'N/A'
                print(f"  {idx}. {source} (Score: {score:.3f})" if isinstance(score, float) else f"  {idx}. {source}")
        
        sources = []
        if hasattr(response, 'source_nodes'):
            for node in response.source_nodes[:3]:
                src = node.metadata.get("source", "Unknown")
                score = node.score if hasattr(node, "score") else None
                sources.append(f"{src} (Score: {score})")

        return response.response, sources
        
    except Exception as e:
        logging.error(f"Query error: {e}")
        print(f"Error processing query: {e}")
        return None

# === Main Function ===

def main():
    """Main execution function"""
    print(" RAG System with LlamaIndex + Groq/Ollama\n")
    
    # Initialize LLM (try Groq first, fallback to Ollama)
    try:
        llm_type = initialize_settings(use_groq=True)
    except:
        llm_type = initialize_settings(use_groq=False)
    
    # Get file paths
    file_input = input("📁 Enter file paths (comma-separated): ").strip()
    file_paths = [f.strip() for f in file_input.split(",") if os.path.exists(f.strip())]
    
    if not file_paths:
        print("No valid file paths provided.")
        return
    
    
    # Create or load index
    index = create_index(file_paths)
    
    if index is None:
        print("Failed to create index.")
        return
    
    print(f"\n Index ready! Using {llm_type.upper()} for responses.")
    print("\nAsk your questions (type 'exit' to quit):\n")
    
    # Query loop
    while True:
        question = input("Question: ").strip()
        if question.lower() in ["exit", "quit", "q"]:
            print("\nGoodbye!")
            break
        
        if not question:
            continue
        
        query_index(index, question, llm_type)
        print()

if __name__ == "__main__":
    main()